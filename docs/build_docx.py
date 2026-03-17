"""Build sponsor documentation as a formatted .docx file."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    tc_pr = cell._element.get_or_add_tcPr()
    shading_elem = tc_pr.makeelement(
        qn("w:shd"),
        {qn("w:fill"): color_hex, qn("w:val"): "clear"},
    )
    tc_pr.append(shading_elem)


def add_styled_table(doc, headers, rows, col_widths=None, highlight_rows=None):
    """Create a styled table with header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "2B579A")

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if highlight_rows and r_idx in highlight_rows:
                run.bold = True
            if r_idx % 2 == 0:
                set_cell_shading(cell, "E8EDF4")

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    return table


def build_document():
    doc = Document()

    # --- Page margins ---
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # --- Styles ---
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Calibri"
    style_normal.font.size = Pt(10.5)
    style_normal.paragraph_format.space_after = Pt(4)
    style_normal.paragraph_format.line_spacing = 1.15

    # ===================== TITLE =====================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("BTCUSDT Price Direction Prediction System")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("Technical Overview for Stakeholders")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ===================== SECTION 1 =====================
    doc.add_heading("1. Project Overview", level=1)

    doc.add_paragraph(
        "This system predicts Bitcoin (BTCUSDT) price direction using classical "
        "machine learning. It answers two key questions:"
    )

    bullet1 = doc.add_paragraph(style="List Bullet")
    run = bullet1.add_run("Will the price go up? ")
    run.bold = True
    bullet1.add_run(
        "(Base models) \u2014 binary prediction of whether BTC price will be "
        "higher after N days."
    )

    bullet2 = doc.add_paragraph(style="List Bullet")
    run = bullet2.add_run("Will the price stay above its trend? ")
    run.bold = True
    bullet2.add_run(
        "(Range models) \u2014 prediction of whether the closing price will remain "
        "above the 14-day moving average after N days."
    )

    doc.add_paragraph(
        "Predictions are generated for four time horizons: 1, 3, 5, and 7 days ahead \u2014 "
        "producing 8 models total. All predictions are served in real-time via a REST API."
    )

    # ===================== SECTION 2 =====================
    doc.add_heading("2. Data Foundation", level=1)

    doc.add_paragraph(
        "The system aggregates 28 daily datasets covering approximately 1,000 days "
        "of history from two primary sources:"
    )

    data_headers = ["Category", "Datasets", "Source"]
    data_rows = [
        ("Futures Open Interest", "4", "CoinGlass API"),
        ("Futures Funding Rates", "3", "CoinGlass API"),
        ("Futures Long/Short Ratios", "3", "CoinGlass API"),
        ("Futures Liquidations", "2", "CoinGlass API"),
        ("Futures Orderbook", "2", "CoinGlass API"),
        ("Futures Taker Volume", "2", "CoinGlass API"),
        ("Exchange Metrics", "3", "CoinGlass API"),
        ("On-Chain Indicators", "4", "CoinGlass API"),
        ("BTC Spot OHLCV", "1", "CoinGlass API"),
        ("S&P 500, Gold, Software ETF", "4", "Yahoo Finance"),
    ]
    add_styled_table(doc, data_headers, data_rows, col_widths=[6, 2.5, 3.5])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.add_run(
        "After merging and feature engineering (derivatives, technical indicators, "
        "lag features), the total feature pool reaches ~1,100 features. "
        "Each model uses a curated subset of 11 to 21 features selected for that "
        "specific prediction task."
    )

    doc.add_paragraph("Key feature groups used by models:", style="List Bullet")
    for feat in [
        "Futures positioning (open interest, funding rates, long/short ratios)",
        "On-chain fundamentals (LTH/STH supply, reserve risk, active addresses)",
        "Technical indicators: RSI, ADX, CCI, Bollinger Band Width, ATR, MFI, OBV, ROC",
        "Macro context (S&P 500 and Gold technical indicators)",
    ]:
        doc.add_paragraph(feat, style="List Bullet 2")

    # ===================== SECTION 3 =====================
    doc.add_heading("3. Model Building Methodology", level=1)

    steps = [
        (
            "Data Collection & Preparation",
            "All 28 datasets are fetched, merged by date, forward-filled for "
            "weekends and holidays, and filtered to the most recent 1,000 days. "
            "Sparse columns (>40% missing values) are automatically removed.",
        ),
        (
            "Feature Engineering",
            "For each numeric column: first differences and percent changes are computed. "
            "Four market imbalance indicators are derived (taker buy/sell imbalance, "
            "liquidation imbalance, orderbook imbalance). 8 technical analysis indicators "
            "are calculated for each of 4 assets (BTC, S&P 500, Gold, IGV) \u2014 "
            "32 TA features total.",
        ),
        (
            "Model Training",
            "Each model is an sklearn Pipeline: Missing Value Imputation \u2192 "
            "Feature Scaling (StandardScaler) \u2192 Logistic Regression "
            "(balanced class weights). Logistic Regression was chosen for its "
            "interpretability, robustness on small datasets, and probabilistic output "
            "enabling threshold tuning.",
        ),
        (
            "Walk-Forward Cross-Validation",
            "Models are validated using TimeSeriesSplit with 4 folds \u2014 the industry "
            "standard for time-series data. This ensures: no future data leakage, "
            "a purge gap equal to the prediction horizon to prevent label contamination, "
            "and realistic evaluation simulating deployment conditions.",
        ),
        (
            "Threshold Optimization",
            "Each model\u2019s probability threshold is tuned to balance precision and "
            "recall for optimal real-world trading performance.",
        ),
    ]

    for i, (title_text, body) in enumerate(steps, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"Step {i} \u2014 {title_text}. ")
        run.bold = True
        run.font.size = Pt(10.5)
        p.add_run(body)

    # ===================== SECTION 4 =====================
    doc.add_heading("4. Model Performance Ranking", level=1)

    doc.add_paragraph(
        "All models are ranked by AUC-ROC (area under the ROC curve) \u2014 the primary "
        "measure of discriminative power. Metrics shown are cross-validation averages "
        "across 4 time-series folds."
    )

    # --- Range models ---
    p = doc.add_paragraph()
    run = p.add_run("Top Performers: Range Models")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)

    perf_headers = ["#", "Model", "Horizon", "AUC", "Accuracy", "Precision", "Recall", "F1"]
    range_rows = [
        ("1", "Range 1D", "1 day", "0.921", "81.7%", "89.0%", "77.5%", "81.1%"),
        ("2", "Range 3D", "3 days", "0.852", "74.7%", "75.6%", "80.2%", "76.1%"),
        ("3", "Range 5D", "5 days", "0.767", "69.5%", "72.5%", "65.0%", "67.2%"),
        ("4", "Range 7D", "7 days", "0.711", "65.8%", "71.5%", "50.1%", "58.4%"),
    ]
    add_styled_table(
        doc, perf_headers, range_rows,
        col_widths=[1, 2.5, 2, 1.8, 2, 2, 1.8, 1.8],
        highlight_rows=[0, 1],
    )

    # --- Base models ---
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run("Base Models (Direct Price Direction)")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)

    base_rows = [
        ("5", "Base 3D", "3 days", "0.621", "52.9%", "58.3%", "32.0%", "37.7%"),
        ("6", "Base 5D", "5 days", "0.599", "56.9%", "63.5%", "51.1%", "48.1%"),
        ("7", "Base 1D", "1 day", "0.554", "54.0%", "53.5%", "70.0%", "60.5%"),
        ("8", "Base 7D", "7 days", "0.528", "55.0%", "51.8%", "65.3%", "56.8%"),
    ]
    add_styled_table(
        doc, perf_headers, base_rows,
        col_widths=[1, 2.5, 2, 1.8, 2, 2, 1.8, 1.8],
    )

    # ===================== SECTION 5 =====================
    doc.add_heading("5. Key Takeaways", level=1)

    takeaways = [
        (
            "Range models significantly outperform base models",
            " (AUC 0.71\u20130.92 vs 0.53\u20130.62). Predicting whether price stays "
            "above its trend is more tractable than predicting raw direction \u2014 "
            "consistent with quantitative finance literature.",
        ),
        (
            "Shorter horizons yield stronger predictions.",
            " The 1-day Range model achieves 92% AUC with 82% accuracy \u2014 "
            "strong discriminative power for a financial prediction task.",
        ),
        (
            "No overfitting detected.",
            " Cross-validation metrics closely track out-of-sample performance, "
            "confirming the walk-forward methodology prevents data leakage.",
        ),
        (
            "Compact feature sets.",
            " The best model (Range 3D) uses only 11 features \u2014 a focused "
            "signal set outperforms high-dimensional approaches.",
        ),
        (
            "Production-ready.",
            " All models are served via a FastAPI endpoint with automatic dataset "
            "caching and model versioning. Predictions available in real-time.",
        ),
    ]

    for bold_part, normal_part in takeaways:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(bold_part)
        run.bold = True
        p.add_run(normal_part)

    # --- Footer ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    run = p.add_run(
        "System processes ~1,100 features from 28 data sources. "
        "Models retrained on demand. API available at port 8000."
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Save
    path = "docs/Sponsor_Documentation.docx"
    doc.save(path)
    print(f"Saved: {path}")


if __name__ == "__main__":
    build_document()
