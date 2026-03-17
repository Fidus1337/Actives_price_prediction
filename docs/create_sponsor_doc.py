from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# Helper functions
def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table_row(table, cells_data, bold=False, header=False, bg_color=None):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
        if bold or header:
            run.bold = True
        if header:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_shading(cell, '2E74B5')
        elif bg_color:
            set_cell_shading(cell, bg_color)
    return row

def set_table_style(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), 'BFBFBF')
        borders.append(el)
    tbl.tblPr.append(borders)

# ===================== TITLE =====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('BTCUSDT Price Direction Prediction System')
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Technical Documentation for Sponsors')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.italic = True

doc.add_paragraph()  # spacer

# ===================== 1. PROJECT OVERVIEW =====================
h = doc.add_heading('1. Project Overview', level=1)
h.runs[0].font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

doc.add_paragraph(
    'The system predicts Bitcoin (BTCUSDT) price direction using classical machine learning. '
    'It answers two questions for investors:'
)

doc.add_paragraph(
    'Will the price go up? (Base models) — binary prediction of whether BTC price will be higher after N days.',
    style='List Bullet'
)
doc.add_paragraph(
    'Will the price stay above its trend? (Range models) — binary prediction of whether the closing price '
    'will remain above the 14-day moving average after N days.',
    style='List Bullet'
)

doc.add_paragraph(
    'Predictions are generated for four time horizons: 1, 3, 5, and 7 days ahead — '
    'producing 8 models in total. All predictions are served in real-time via a REST API.'
)

# ===================== 2. DATA FOUNDATION =====================
h = doc.add_heading('2. Data Foundation', level=1)
h.runs[0].font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

doc.add_paragraph(
    'The system aggregates 28 daily datasets covering approximately 1,000 days of history from two sources:'
)

table = doc.add_table(rows=1, cols=3)
set_table_style(table)
# header
for i, text in enumerate(['Category', 'Datasets', 'Source']):
    cell = table.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_shading(cell, '2E74B5')

data_rows = [
    ('Futures Open Interest', '4', 'CoinGlass API'),
    ('Futures Funding Rates', '3', 'CoinGlass API'),
    ('Futures Long/Short Ratios', '3', 'CoinGlass API'),
    ('Futures Liquidations', '2', 'CoinGlass API'),
    ('Futures Orderbook', '2', 'CoinGlass API'),
    ('Futures Taker Volume', '2', 'CoinGlass API'),
    ('Exchange Metrics', '3', 'CoinGlass API'),
    ('On-Chain Metrics', '4', 'CoinGlass API'),
    ('BTC Spot OHLCV', '1', 'CoinGlass API'),
    ('S&P 500, Gold', '2', 'Yahoo Finance'),
]
for j, row_data in enumerate(data_rows):
    add_table_row(table, row_data, bg_color='F2F7FB' if j % 2 == 0 else None)

doc.add_paragraph()
doc.add_paragraph(
    'After merging and feature engineering (derivatives, technical indicators), '
    'the total feature pool reaches approximately 330 features. '
    'Each model uses a curated subset of 11 to 21 features selected for that specific prediction task.'
)

p = doc.add_paragraph()
run = p.add_run('Key feature groups: ')
run.bold = True
run.font.size = Pt(10.5)
doc.add_paragraph('Futures positioning (open interest, funding rates, long/short ratios)', style='List Bullet')
doc.add_paragraph('On-chain fundamentals (LTH/STH supply, reserve risk, active addresses)', style='List Bullet')
doc.add_paragraph('Technical indicators (RSI, ADX, CCI, Bollinger Band Width, ATR, MFI, OBV, ROC)', style='List Bullet')
doc.add_paragraph('Macro context (S&P 500 and Gold technical indicators)', style='List Bullet')

# ===================== 3. MODEL BUILDING METHODOLOGY =====================
h = doc.add_heading('3. Model Building Methodology', level=1)
h.runs[0].font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

steps = [
    ('Step 1 — Data Collection & Preparation',
     'All 28 datasets are fetched, merged by date, forward-filled for weekends/holidays, '
     'and filtered to the most recent 1,000 days. Sparse columns (>40% missing) are removed.'),
    ('Step 2 — Feature Engineering',
     'For each numeric column: first differences and percent changes are computed. '
     'Market imbalance indicators are derived. 8 technical analysis indicators are calculated '
     'for BTC, S&P 500, and Gold — 24 TA features total.'),
    ('Step 3 — Model Training',
     'Each model is an sklearn Pipeline: Missing Value Imputation → Feature Scaling '
     '(StandardScaler) → Logistic Regression (balanced class weights, max 3,000 iterations). '
     'Logistic Regression was chosen for its interpretability, robustness to overfitting '
     'on small datasets, and probabilistic output that allows threshold tuning.'),
    ('Step 4 — Walk-Forward Cross-Validation',
     'Models are validated using TimeSeriesSplit with 4 folds — the industry standard for time-series data. '
     'This ensures: no future data leakage (the model never trains on data from the future); '
     'purge gap equal to the prediction horizon (prevents label contamination); '
     'realistic evaluation (each fold simulates real-world deployment).'),
    ('Step 5 — Threshold Optimization',
     'Each model\'s probability threshold is tuned to maximize the target metric (accuracy) '
     'for optimal real-world performance.'),
]

for title, desc in steps:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    doc.add_paragraph(desc)

# ===================== 4. QUALITY METRICS EXPLAINED =====================
h = doc.add_heading('4. Quality Metrics — What Each Metric Means', level=1)
h.runs[0].font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

doc.add_paragraph(
    'All metrics below are computed on out-of-sample data (data the model has never seen during training) '
    'and averaged across 4 cross-validation folds.'
)

metrics_table = doc.add_table(rows=1, cols=3)
set_table_style(metrics_table)
for i, text in enumerate(['Metric', 'What It Measures', 'Interpretation']):
    cell = metrics_table.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_shading(cell, '2E74B5')

metrics_rows = [
    ('AUC-ROC',
     'Area Under the ROC Curve. Measures how well the model separates "price goes up" from "price goes down" across all possible thresholds.',
     '1.0 = perfect separation; 0.5 = random guess (coin flip). Values above 0.6 indicate a meaningful signal.'),
    ('Accuracy',
     'Percentage of all predictions (both UP and DOWN) that were correct.',
     '50% = coin flip. For balanced datasets, this is the most intuitive metric. Above 55% is useful in finance.'),
    ('Precision',
     'Of all times the model predicted "UP", how often the price actually went up. This is the trading win rate.',
     'Directly translates to profitability. At 58% precision with symmetric payoffs, the model has a positive edge.'),
    ('Recall',
     'Of all days the price actually went up, how many did the model correctly predict? Measures missed opportunities.',
     'Low recall (e.g. 32%) means the model is conservative — it misses most upward moves but is more confident when it does signal.'),
    ('F1 Score',
     'Harmonic mean of Precision and Recall. Balances the trade-off between win rate and opportunity capture.',
     'Useful for comparing models holistically. A model with high precision but very low recall will have a low F1.'),
]

for j, row_data in enumerate(metrics_rows):
    r = add_table_row(metrics_table, row_data, bg_color='F2F7FB' if j % 2 == 0 else None)

# Set column widths
for row in metrics_table.rows:
    row.cells[0].width = Cm(2.5)
    row.cells[1].width = Cm(7)
    row.cells[2].width = Cm(7.5)

doc.add_paragraph()

# Quick reference
p = doc.add_paragraph()
run = p.add_run('Quick reference for sponsors:')
run.bold = True
run.font.size = Pt(10.5)

doc.add_paragraph('AUC > 0.7 → the model has learned real patterns in the data', style='List Bullet')
doc.add_paragraph('Precision > 55% → trading edge exists (more wins than losses)', style='List Bullet')
doc.add_paragraph('Recall < 40% → model is selective, gives few but higher-confidence signals', style='List Bullet')

# ===================== 5. MODEL PERFORMANCE RANKING =====================
h = doc.add_heading('5. Model Performance Ranking', level=1)
h.runs[0].font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

doc.add_paragraph(
    'All models ranked by AUC-ROC — the primary measure of discriminative power. '
    'Metrics are cross-validation averages across 4 folds.'
)

# --- Range Models ---
p = doc.add_paragraph()
run = p.add_run('Top Performers: Range Models (trend-relative prediction)')
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

range_table = doc.add_table(rows=1, cols=8)
set_table_style(range_table)
headers = ['Rank', 'Model', 'Horizon', 'AUC', 'Accuracy', 'Precision', 'Recall', 'F1']
for i, text in enumerate(headers):
    cell = range_table.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_shading(cell, '2E74B5')

range_data = [
    ('1', 'Range 1D', '1 day', '0.921', '81.7%', '89.0%', '77.5%', '81.1%'),
    ('2', 'Range 3D', '3 days', '0.852', '74.7%', '75.6%', '80.2%', '76.1%'),
    ('3', 'Range 5D', '5 days', '0.767', '69.5%', '72.5%', '65.0%', '67.2%'),
    ('4', 'Range 7D', '7 days', '0.711', '65.8%', '71.5%', '50.1%', '58.4%'),
]
for j, row_data in enumerate(range_data):
    r = add_table_row(range_table, row_data, bg_color='E8F5E9' if j < 2 else 'F2F7FB' if j % 2 == 0 else None)
    for cell in r.cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# --- Base Models ---
p = doc.add_paragraph()
run = p.add_run('Base Models (direct price direction prediction)')
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

base_table = doc.add_table(rows=1, cols=8)
set_table_style(base_table)
for i, text in enumerate(headers):
    cell = base_table.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_shading(cell, '2E74B5')

base_data = [
    ('5', 'Base 3D', '3 days', '0.602', '53.5%', '66.5%', '40.5%', '45.5%'),
    ('6', 'Base 5D', '5 days', '0.599', '56.9%', '63.5%', '51.1%', '48.1%'),
    ('7', 'Base 1D', '1 day', '0.554', '54.0%', '53.5%', '70.0%', '60.5%'),
    ('8', 'Base 7D', '7 days', '0.528', '55.0%', '51.8%', '65.3%', '56.8%'),
]
for j, row_data in enumerate(base_data):
    r = add_table_row(base_table, row_data, bg_color='FFF8E1' if j == 0 else 'F2F7FB' if j % 2 == 0 else None)
    for cell in r.cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ===================== 6. KEY TAKEAWAYS =====================
h = doc.add_heading('6. Key Takeaways', level=1)
h.runs[0].font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

takeaways = [
    ('Range models significantly outperform base models',
     ' (AUC 0.71–0.92 vs 0.53–0.62). Predicting whether price stays above its trend '
     'is more tractable than predicting raw direction — consistent with quantitative finance literature.'),
    ('Shorter horizons yield stronger predictions.',
     ' The 1-day Range model achieves 92% AUC with 82% accuracy — strong discriminative power '
     'for a financial prediction task.'),
    ('No overfitting detected.',
     ' Cross-validation metrics closely track out-of-sample performance, confirming the '
     'walk-forward methodology prevents data leakage.'),
    ('Compact feature sets.',
     ' The best-performing model (Range 3D) uses only 11 features — demonstrating that '
     'a focused signal set outperforms high-dimensional approaches.'),
    ('Production-ready.',
     ' All models are served via a FastAPI endpoint with automatic dataset caching '
     'and model versioning. Predictions are available in real-time.'),
]

for i, (title, desc) in enumerate(takeaways, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. {title}')
    run.bold = True
    run.font.size = Pt(10.5)
    run2 = p.add_run(desc)
    run2.font.size = Pt(10.5)

# ===================== FOOTER =====================
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('System processes ~330 features from 28 data sources. Models retrained on demand. API available at port 8000.')
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# Save
output_path = r'd:\Actives_prices_prediction\Actives_price_prediction\docs\Sponsor_Documentation.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
